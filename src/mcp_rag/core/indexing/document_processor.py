"""Document processing for the indexing foundation."""

from __future__ import annotations

import hashlib
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Optional

from .models import ChunkRecord, ProcessedDocument, IndexingSettings
from .text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process raw text and supported files into chunk records."""

    def __init__(self, settings: IndexingSettings | None = None):
        self.settings = settings or IndexingSettings()
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.settings.chunk_size,
            chunk_overlap=self.settings.chunk_overlap,
            separators=self.settings.separators,
        )

    def process_text(
        self,
        text: str,
        *,
        source: str = "inline_text",
        filename: str | None = None,
        file_type: str = "text",
        metadata: Dict[str, Any] | None = None,
    ) -> ProcessedDocument:
        """Process raw text into a normalized document payload."""

        resolved_filename = filename or source or "inline_text"
        resolved_source = source or resolved_filename
        content_hash = hashlib.blake2b(text.encode("utf-8"), digest_size=16).hexdigest()
        document_id = self._build_document_id(resolved_source, content_hash)

        document_metadata = self._build_metadata(
            source=resolved_source,
            filename=resolved_filename,
            file_type=file_type,
            content=text,
            extra_metadata=metadata,
        )
        document_metadata["content_hash"] = content_hash

        return ProcessedDocument(
            document_id=document_id,
            source=resolved_source,
            filename=resolved_filename,
            file_type=file_type,
            content=text,
            metadata=document_metadata,
        )

    def process_file(
        self,
        file_path: str | Path,
        *,
        metadata: Dict[str, Any] | None = None,
        filename: str | None = None,
    ) -> ProcessedDocument:
        """Process a supported file into a normalized document payload."""

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        suffix = path.suffix.lower()
        file_type = suffix.lstrip(".") or "text"
        resolved_filename = filename or path.name
        source = str(path)

        try:
            if suffix in {".txt", ".md", ".text", ""}:
                content, encoding = self._load_text(path)
            elif suffix == ".pdf":
                content, encoding = self._load_pdf(path)
            elif suffix == ".docx":
                content, encoding = self._load_docx(path)
            else:
                content, encoding = self._load_text(path)
        except Exception as exc:
            logger.error("Failed to process file %s: %s", path, exc)
            return ProcessedDocument(
                document_id=self._build_document_id(source, "error"),
                source=source,
                filename=resolved_filename,
                file_type=file_type,
                content="",
                metadata={
                    "source": source,
                    "filename": resolved_filename,
                    "file_type": file_type,
                    "schema_version": self.settings.schema_version,
                    "processed_at": datetime.now(timezone.utc).isoformat(),
                    "error": str(exc),
                },
                error=str(exc),
            )

        file_stat = path.stat()
        extra = {
            "file_size": file_stat.st_size,
            "modified_time": datetime.fromtimestamp(file_stat.st_mtime, tz=timezone.utc).isoformat(),
            "encoding": encoding,
        }
        if metadata:
            extra.update(metadata)

        return self.process_text(
            content,
            source=source,
            filename=resolved_filename,
            file_type=file_type,
            metadata=extra,
        )

    def process_text_to_chunks(
        self,
        text: str,
        *,
        source: str = "inline_text",
        filename: str | None = None,
        file_type: str = "text",
        metadata: Dict[str, Any] | None = None,
    ) -> list[ChunkRecord]:
        return self.chunk_document(
            self.process_text(
                text,
                source=source,
                filename=filename,
                file_type=file_type,
                metadata=metadata,
            )
        )

    def process_file_to_chunks(
        self,
        file_path: str | Path,
        *,
        metadata: Dict[str, Any] | None = None,
        filename: str | None = None,
    ) -> list[ChunkRecord]:
        return self.chunk_document(
            self.process_file(file_path, metadata=metadata, filename=filename)
        )

    def chunk_document(
        self,
        document: ProcessedDocument,
        *,
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
    ) -> list[ChunkRecord]:
        """Split a processed document into chunk records."""

        if document.error:
            return []

        splitter = self.splitter
        if chunk_size is not None or chunk_overlap is not None:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size or self.settings.chunk_size,
                chunk_overlap=chunk_overlap if chunk_overlap is not None else self.settings.chunk_overlap,
                separators=self.settings.separators,
            )

        chunk_texts = splitter.split_text(document.content)
        total_chunks = len(chunk_texts)
        chunks: list[ChunkRecord] = []

        for index, chunk_text in enumerate(chunk_texts):
            chunk_id = f"{document.document_id}_chunk_{index:04d}"
            chunk_metadata = {
                **document.metadata,
                "schema_version": self.settings.schema_version,
                "document_id": document.document_id,
                "chunk_id": chunk_id,
                "chunk_index": index,
                "total_chunks": total_chunks,
                "source": document.source,
                "filename": document.filename,
                "file_type": document.file_type,
                "chunk_char_count": len(chunk_text),
                "chunk_line_count": chunk_text.count("\n") + 1 if chunk_text else 0,
            }
            chunks.append(
                ChunkRecord(
                    chunk_id=chunk_id,
                    document_id=document.document_id,
                    chunk_index=index,
                    total_chunks=total_chunks,
                    source=document.source,
                    filename=document.filename,
                    file_type=document.file_type,
                    content=chunk_text,
                    metadata=chunk_metadata,
                )
            )

        return chunks

    def _build_metadata(
        self,
        *,
        source: str,
        filename: str,
        file_type: str,
        content: str,
        extra_metadata: Dict[str, Any] | None = None,
    ) -> Dict[str, Any]:
        metadata = {
            "schema_version": self.settings.schema_version,
            "source": source,
            "filename": filename,
            "file_type": file_type,
            "char_count": len(content),
            "line_count": content.count("\n") + 1 if content else 0,
            "word_count": len(content.split()),
            "processed_at": datetime.now(timezone.utc).isoformat(),
        }
        if extra_metadata:
            metadata.update(extra_metadata)
        return metadata

    def _build_document_id(self, source: str, content_hash: str) -> str:
        payload = f"{source}|{content_hash}"
        return hashlib.blake2b(payload.encode("utf-8"), digest_size=16).hexdigest()

    def _load_text(self, path: Path) -> tuple[str, str]:
        encodings = ("utf-8", "utf-8-sig", "gbk", "gb2312", "latin1")
        last_error: Exception | None = None
        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding), encoding
            except UnicodeDecodeError as exc:
                last_error = exc
        raise ValueError(f"Unable to decode text file {path}: {last_error}")

    def _load_pdf(self, path: Path) -> tuple[str, str]:
        try:
            import PyPDF2
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ImportError("PyPDF2 is required for PDF processing") from exc

        with path.open("rb") as handle:
            reader = PyPDF2.PdfReader(handle)
            pages: list[str] = []
            for page in reader.pages:
                pages.append(page.extract_text() or "")
        return "\n".join(pages), "binary-pdf"

    def _load_docx(self, path: Path) -> tuple[str, str]:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ImportError("python-docx is required for DOCX processing") from exc

        doc = Document(path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    paragraphs.append(row_text)
        return "\n".join(paragraphs), "docx"

